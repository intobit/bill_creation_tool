import os
from docx import Document
from docx.shared import Cm, Pt
from datetime import datetime
from wordReader import WordReader


class CreateWordDoc(WordReader):
    def __init__(self, document: str, directory_path: str):
        super().__init__(document)
        self.doc = Document()
        self.output_path = directory_path
        self.output_file_name = "testl123.docx"
        self.font_name = 'Montserrat'
        self.font_size_6 = 6
        self.font_size_8 = 8
        self.font_size_9 = 9
        self.font_size_10 = 10

    def create_bill(self):
        self._set_page_layout()
        self._header_section()
        self._add_company_data()
        self._invoice_num_delivery_date()
        self._acknowledgement()
        table, columns, leftover_rows = self._invoice_table()
        table = self._table_column_width(table)
        self._table_styling(table, columns, leftover_rows)
        self._payment_data()
        self._signature()
        self._regards()
        self._footer_section()
        self._generate_document()

    def _set_page_layout(self, top: float = 0, bottom: float = 2.54, left: float = 2.54, right: float = 1.5):
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Cm(top)
            section.bottom_margin = Cm(bottom)
            section.left_margin = Cm(left)
            section.right_margin = Cm(right)

    def _header_section(self, header_distance: int = 0, empty_line: str = ""):
        section = self.doc.sections[0]
        header = section.header
        paragraph = header.paragraphs[0]
        # Set the header distance to 0 to remove top and bottom margins
        section.header_distance = Pt(header_distance)

        run = paragraph.add_run()
        run.add_picture('./images/header_logo.png')
        self.doc.add_paragraph(empty_line)

    def _add_formatted_paragraph(self, content, font_name, font_size):
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run(content)
        run.font.name = font_name
        run.font.size = Pt(font_size)

    def _add_company_data(self):
        company_name, person_name, street_address, city = self.get_company_data()
        actual_date = datetime.now().date()
        format_date = actual_date.strftime('%d.%m.%Y')
        content = (f'{company_name}\t\t\t\t\t\t\t\t\t\t{format_date}\n'
                   f'{person_name}\n'
                   f'{street_address}\n'
                   f'{city}')
        self._add_formatted_paragraph(content, self.font_name, self.font_size_9)

    def _invoice_num_delivery_date(self):
        actual_date = datetime.now().date()
        format_date = actual_date.strftime('%d.%m.%Y')
        content = f'Rechnung Nr. xx\nLieferdatum: {format_date}'
        self._add_formatted_paragraph(content, self.font_name, self.font_size_10)

    def _acknowledgement(self):
        content = 'Vielen Dank für deine Bestellung bei Company XYZ! Hiermit stellen wir folgende Rechnung:'
        self._add_formatted_paragraph(content, self.font_name, self.font_size_9)

    def _invoice_table(self, fixed_rows: int = 4, columns: int = 4, brutto_price: float = 4.2,
                       shipping_cost: float = 4.75):
        items_dict = self.create_dict_from_table()

        total_rows = fixed_rows + len(items_dict)
        table = self.doc.add_table(rows=total_rows, cols=columns)

        header_row = table.rows[0]
        header_row.cells[0].text = 'Produktname'
        header_row.cells[1].text = 'Einzelpreis (brutto)'
        header_row.cells[2].text = 'Anzahl'
        header_row.cells[3].text = 'Gesamt'

        for i in range(1, len(items_dict) + 1):
            item_row = table.rows[i]
            item_row.cells[0].text = f'{list(items_dict)[i - 1]}'
            txt_brutto_price = '{:.2f}'.format(brutto_price).replace('.', ',')
            item_row.cells[1].text = f'{txt_brutto_price}€'
            item_row.cells[2].text = f'{list(items_dict.values())[i - 1]}'
            txt_total_price = '{:.2f}'.format(int(list(items_dict.values())[i - 1]) * brutto_price).replace('.', ',')
            item_row.cells[3].text = f'{txt_total_price}€'

        leftover_rows = len(items_dict) + 1

        shipping_row = table.rows[leftover_rows]
        shipping_row.cells[0].text = 'Versand'
        shipping_row.cells[1].text = ''
        shipping_row.cells[2].text = ''
        txt_shipping_cost = '{:.2f}'.format(shipping_cost).replace('.', ',')
        shipping_row.cells[3].text = f'{txt_shipping_cost}€'

        total_amount_row = table.rows[leftover_rows + 1]
        total_amount_row.cells[0].text = 'Gesamtbetrag'
        total_amount_row.cells[1].text = ''
        total_amount_row.cells[2].text = ''
        txt_total_amount = '{:.2f}'.format(self.total_price() + shipping_cost).replace('.', ',')
        total_amount_row.cells[3].text = f'{txt_total_amount}€'

        netto_row = table.rows[leftover_rows + 2]
        netto_row.cells[0].text = 'Netto'
        netto_row.cells[1].text = ''
        netto_row.cells[2].text = ''
        txt_netto = '{:.2f}'.format((self.total_price() + shipping_cost)/1.2).replace('.', ',')
        netto_row.cells[3].text = f'{txt_netto}€'

        return table, columns, leftover_rows

    @staticmethod
    def _table_column_width(table):
        col_widths = [10.5, 3.28, 1.16, 1.98]

        i = 0
        while i <= 3:
            for cell in table.columns[i].cells:
                cell.width = Cm(col_widths[i])
            i += 1
        return table

    @staticmethod
    def _table_styling(table, columns, leftover_rows, font_name: str = 'Montserrat', font_size: int = 8,
                       space_after: int = 2, line_spacing: float = 1, bold: bool = True):
        i = 0
        while i <= columns - 1:
            row_count = 0
            for cell in table.columns[i].cells:
                paragraphs = cell.paragraphs

                for paragraph in paragraphs:
                    paragraph.paragraph_format.space_after = Pt(space_after)
                    paragraph.paragraph_format.line_spacing = line_spacing
                    if i == 0:
                        paragraph.alignment = 0     # Text left
                    if i > 0:
                        paragraph.alignment = 1     # Text center

                    for run in paragraph.runs:
                        font = run.font
                        font.name = font_name
                        if i < 4 and 0 < row_count < leftover_rows:
                            font.size = Pt(6)
                        elif cell.text == 'Gesamtbetrag':
                            font.size = Pt(10)
                            run.bold = bold
                        elif cell.text == 'Gesamt':
                            run.bold = bold
                        elif i == 3 and row_count == leftover_rows + 1:
                            font.size = Pt(9)
                            run.bold = bold
                        else:
                            font.size = Pt(font_size)
                row_count += 1
            i += 1

        # Change font size of first row back to 8
        for cell in table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(8)
        return table

    def _payment_data(self):
        content = ('Zahlbar innerhalb von 14 Tagen per Überweisung an:\n'
                   'Company XYZ \n'
                   'IBAN: AT12 3456 7891 2345 6789\n'
                   'BIC: ABCDEFGHIJKL\n\n'
                   'Für Rückfragen stehen wir gerne zur Verfügung\n'
                   'Liebe Grüße')
        self._add_formatted_paragraph(content, self.font_name, self.font_size_9)

    def _signature(self):
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run()
        run.add_picture('./images/signature.png')

    def _regards(self):
        content = 'Klara Müller\nfür das Company XYZ-Team'
        self._add_formatted_paragraph(content, self.font_name, self.font_size_9)

    def _footer_section(self):
        section = self.doc.sections[0]
        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = 1     # Text center
        run = paragraph.add_run(
            'Company XYZ; FN: 11111f; FB-Gericht: Handelsgericht Wien; Sitz: 1100 Wien; UID: ATU123456')
        run.font.name = self.font_name
        run.font.size = Pt(self.font_size_8)

    def _generate_document(self):
        directory_path = self.output_path
        filename = self.output_file_name
        file_path = os.path.join(directory_path, filename)
        if os.path.isfile(file_path):
            print(f"{filename} exists in {directory_path}. Deleting it...")
            os.remove(file_path)
        self.doc.save(file_path)
        print(f"Created new {filename}...in {file_path}")
