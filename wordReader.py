from docx import Document


class WordReader:
    def __init__(self, document: str):
        self.document = document

    def get_company_data(self):
        doc = Document(self.document)
        line_counter = 0
        company_name = ""
        person_name = ""
        street_address = ""
        city = ""
        for paragraph in doc.paragraphs:
            line_counter += 1
            if line_counter == 2:
                company_name = paragraph.text
            if line_counter == 3:
                person_name = paragraph.text
            if line_counter == 4:
                street_address = paragraph.text
            if line_counter == 5:
                city = paragraph.text
        return company_name, person_name, street_address, city

    def get_table(self):
        doc = Document(self.document)
        item_amount_lst = []
        row_counter = 0
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if row_counter > 3 and not cell.text.isupper() and cell.text != "Produktname" \
                            and cell.text != "Anzahl":
                        item_amount_lst.append(cell.text)
                    row_counter += 1
        return item_amount_lst

    def create_dict_from_table(self):
        data_dict = {}
        table = self.get_table()
        for i in range(0, len(table), 2):
            item = table[i]
            if table[i + 1] == '':
                continue
            else:
                amount = table[i + 1]
                int(amount)
            data_dict[item] = amount
        return data_dict

    def total_price(self, price: float = 4.2):
        dict = self.create_dict_from_table()
        total = 0
        price = price
        for value in dict.values():
            total += int(value)
        final_price = round(total * price, 2)
        return final_price
